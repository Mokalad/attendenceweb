# Streamlit Web App to Analyze PDF Attendance and Export Word Report

import streamlit as st
import pandas as pd
import pdfplumber
from docx import Document
from datetime import datetime, timedelta
import calendar
from io import BytesIO

# --- Functions ---
def is_morning_shift(arrival_time): return 9 <= arrival_time.hour < 14
def is_evening_shift(arrival_time): return arrival_time.hour >= 14
def is_double_shift(arrival_time, departure_time): return arrival_time.hour < 14 and departure_time.hour >= 22
def is_single_punch_shift(punch_times): return len(punch_times) == 1
def count_shifts(shift_type): return 1 if shift_type in ['مسائية', 'صباحية', 'بصمة واحدة'] else 2 if shift_type == 'مزدوجة' else 0
def calculate_overtime(total): return total - 26 if total > 26 else 0
def count_delays(punch_times): return len([p for p in punch_times if (p.hour == 15 and p.minute > 10) or (p.hour == 16)]), [p for p in punch_times if (p.hour == 15 and p.minute > 10) or (p.hour == 16)]

arabic_to_english_map = str.maketrans('٠١٢٣٤٥٦٧٨٩', '0123456789')
def convert_arabic_numerals(text): return text.translate(arabic_to_english_map) if isinstance(text, str) else text
def format_time_12h(time_obj): return time_obj.strftime("%I:%M %p")

def load_pdf_data(uploaded_file):
    data = []
    with pdfplumber.open(uploaded_file) as pdf:
        for page in pdf.pages:
            table = page.extract_table()
            if table and len(table) > 1:
                data.extend(table[1:])
    return data

def process_attendance_data(data):
    if not data:
        return pd.DataFrame()

    df = pd.DataFrame(data, columns=["رقم البصمه", "الإسم", "التاريخ_الوقت"])
    df['cleaned'] = df['التاريخ_الوقت'].apply(convert_arabic_numerals)
    df['formatted'] = df['cleaned'].astype(str).str.replace('م', 'PM').str.replace('ص', 'AM')
    df['datetime'] = pd.to_datetime(df['formatted'], format="%d/%m/%Y %I:%M %p", errors='coerce')
    df.dropna(subset=['datetime'], inplace=True)

    min_date = df['datetime'].min().date()
    max_date = df['datetime'].max().date()
    all_days = [min_date + timedelta(days=x) for x in range((max_date - min_date).days + 1)]

    summary = []

    for name, group in df.groupby("الإسم"):
        group = group.sort_values('datetime')
        shifts_count = 0
        delays_count = 0
        shift_types = set()
        delay_details = []
        morning_shift_details = []
        late_departure_details = []
        attended_dates = set()
        double_shift_details = []

        for day, punches in group.groupby(group['datetime'].dt.date):
            punch_times = punches['datetime'].tolist()
            if not punch_times:
                continue

            attended_dates.add(day)
            arrival = min(punch_times)
            departure = max(punch_times)

            if is_single_punch_shift(punch_times):
                shift_type = 'بصمة واحدة'
            elif is_double_shift(arrival, departure):
                shift_type = 'مزدوجة'
                duration = departure - arrival
                double_shift_details.append([str(day), format_time_12h(arrival), format_time_12h(departure), str(duration)])
            elif is_morning_shift(arrival):
                shift_type = 'صباحية'
                morning_shift_details.append([str(day), format_time_12h(arrival), format_time_12h(departure)])
            elif is_evening_shift(arrival):
                shift_type = 'مسائية'
            else:
                shift_type = 'غير معروف'

            shifts_count += count_shifts(shift_type)
            shift_types.add(shift_type)
            daily_delays, delay_times = count_delays(punch_times)
            delays_count += daily_delays
            for delay_time in delay_times:
                delay_details.append([str(day), format_time_12h(delay_time)])

            if departure.hour > 22 or (departure.hour == 22 and departure.minute >= 20):
                late_departure_details.append([str(day), format_time_12h(departure)])

        overtime = calculate_overtime(shifts_count)
        absent_days = [d for d in all_days if d not in attended_dates]
        absent_days_formatted = [[f"{day.strftime('%Y-%m-%d')} ({calendar.day_name[day.weekday()]})"] for day in absent_days]

        summary.append({
            "الإسم": name,
            "عدد الورديات": shifts_count,
            "نوع الورديات": ', '.join(shift_types),
            "عدد التأخيرات": delays_count,
            "تفاصيل التأخيرات": delay_details,
            "أيام الغياب": len(absent_days),
            "تفاصيل أيام الغياب": absent_days_formatted,
            "الدوام الإضافي": overtime,
            "تفاصيل الورديات الصباحية": morning_shift_details,
            "تفاصيل الخروج المتأخر": late_departure_details,
            "تفاصيل الورديات المزدوجة": double_shift_details
        })

    return pd.DataFrame(summary)

def export_to_word(summary_df):
    buffer = BytesIO()
    doc = Document()
    doc.add_heading("تقرير الحضور والانصراف العام", level=0)
    doc.add_paragraph("---")
    doc.add_paragraph()

    for _, row in summary_df.iterrows():
        doc.add_heading(f"بيانات الموظف: {row['الإسم']}", level=1)
        p = doc.add_paragraph()
        p.add_run(f"• عدد الورديات: {row['عدد الورديات']}\n")
        p.add_run(f"• نوع الورديات: {row['نوع الورديات']}\n")
        p.add_run(f"• عدد التأخيرات: {row['عدد التأخيرات']}\n")
        p.add_run(f"• أيام الغياب: {row['أيام الغياب']}\n")
        p.add_run(f"• الدوام الإضافي: {row['الدوام الإضافي']}\n")

        if row["تفاصيل الورديات المزدوجة"]:
            doc.add_paragraph("تفاصيل الورديات المزدوجة:").bold = True
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'التاريخ'
            hdr_cells[1].text = 'وقت الحضور'
            hdr_cells[2].text = 'وقت الانصراف'
            hdr_cells[3].text = 'المدة'
            for d in row["تفاصيل الورديات المزدوجة"]:
                cells = table.add_row().cells
                cells[0].text, cells[1].text, cells[2].text, cells[3].text = d

        doc.add_paragraph("=" * 60)
        doc.add_paragraph()

    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- Streamlit Interface ---
st.set_page_config(page_title="تحليل الحضور", layout="centered")
st.title("📄 أداة تحليل الحضور وإنشاء تقرير Word")
uploaded_file = st.file_uploader("ارفع ملف PDF لبيانات الحضور", type="pdf")

if uploaded_file:
    with st.spinner("🔍 جاري تحليل البيانات..."):
        data = load_pdf_data(uploaded_file)
        summary_df = process_attendance_data(data)

    if summary_df.empty:
        st.warning("⚠️ لا توجد بيانات صالحة بعد التحليل.")
    else:
        st.success("✅ تم التحليل بنجاح! يمكنك تحميل التقرير أدناه.")
        word_buffer = export_to_word(summary_df)
        st.download_button(
            label="📥 تحميل تقرير Word",
            data=word_buffer,
            file_name="تقرير_الحضور.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
