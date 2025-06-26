
import streamlit as st
import pandas as pd
import pdfplumber
from docx import Document
from datetime import datetime, timedelta
from io import BytesIO
import calendar

def is_morning_shift(arrival_time):
    return 9 <= arrival_time.hour < 14

def is_evening_shift(arrival_time):
    return arrival_time.hour >= 14

def is_double_shift(arrival_time, departure_time):
    return arrival_time.hour < 14 and departure_time.hour >= 22

def is_single_punch_shift(punch_times):
    return len(punch_times) == 1

def count_shifts(shift_type):
    if shift_type in ['مسائية', 'صباحية', 'بصمة واحدة']:
        return 1
    elif shift_type == 'مزدوجة':
        return 2
    return 0

def calculate_overtime(total_shifts):
    return total_shifts - 26 if total_shifts > 26 else 0

def count_delays(punch_times):
    delays = [p for p in punch_times if (p.hour == 15 and p.minute > 10) or (p.hour == 16 and p.minute < 60)]
    return len(delays), delays

arabic_to_english_map = str.maketrans('٠١٢٣٤٥٦٧٨٩','0123456789')
def convert_arabic_numerals(text):
    if isinstance(text, str):
        return text.translate(arabic_to_english_map)
    return text

def format_time_12h(time_obj):
    return time_obj.strftime("%I:%M %p")

def load_pdf_data(pdf_file):
    data = []
    with pdfplumber.open(pdf_file) as pdf:
        for page in pdf.pages:
            table = page.extract_table()
            if table and len(table) > 1:
                data.extend(table[1:])
    return data

def process_attendance_data(data):
    df = pd.DataFrame(data, columns=["رقم البصمه", "الإسم", "التاريخ_الوقت"])
    df['التاريخ_الوقت_cleaned'] = df['التاريخ_الوقت'].apply(convert_arabic_numerals)
    df['التاريخ_الوقت_final'] = df['التاريخ_الوقت_cleaned'].astype(str).str.replace('م', 'PM').str.replace('ص', 'AM')
    df['التاريخ_الوقت'] = pd.to_datetime(df['التاريخ_الوقت_final'], format="%d/%m/%Y %I:%M %p", errors='coerce')
    df.dropna(subset=['التاريخ_الوقت'], inplace=True)

    if df.empty:
        return pd.DataFrame()

    min_date = df['التاريخ_الوقت'].min().date()
    max_date = df['التاريخ_الوقت'].max().date()
    all_dates_in_period = [min_date + timedelta(days=x) for x in range((max_date - min_date).days + 1)]

    summary = []

    for name, group in df.groupby("الإسم"):
        group = group.sort_values('التاريخ_الوقت')
        shifts_count = 0
        delays_count = 0
        shift_types = set()
        delay_details = []
        morning_shift_details = []
        late_departure_details = []
        attended_dates = set()
        double_shift_details = []

        for day, punches in group.groupby(group['التاريخ_الوقت'].dt.date):
            attended_dates.add(day)
            punch_times = punches['التاريخ_الوقت'].tolist()
            arrival = min(punch_times)
            departure = max(punch_times)

            shift_type = None
            if is_single_punch_shift(punch_times):
                shift_type = 'بصمة واحدة'
            elif is_double_shift(arrival, departure):
                shift_type = 'مزدوجة'
                duration = departure - arrival
                double_shift_details.append([str(day), format_time_12h(arrival), format_time_12h(departure), str(duration)])
            elif is_morning_shift(arrival):
                shift_type = 'صباحية'
                morning_shift_details.append([str(day), format_time_12h(arrival), format_time_12h(departure)])
            elif is_evening_shift(arrival):
                shift_type = 'مسائية'

            if shift_type:
                shifts_count += count_shifts(shift_type)
                shift_types.add(shift_type)

            daily_delays, delay_times_list = count_delays(punch_times)
            delays_count += daily_delays
            for delay_time in delay_times_list:
                delay_details.append([str(day), format_time_12h(delay_time)])

            if departure.hour > 22 or (departure.hour == 22 and departure.minute >= 20):
                late_departure_details.append([str(day), format_time_12h(departure)])

        overtime_count = calculate_overtime(shifts_count)
        absent_days = sorted([d for d in all_dates_in_period if d not in attended_dates])
        absent_days_formatted = [[f"{day.strftime('%Y-%m-%d')} ({calendar.day_name[day.weekday()]})"] for day in absent_days]

        summary.append({
            "الإسم": name,
            "عدد الورديات": shifts_count,
            "نوع الورديات": ', '.join(shift_types) if shift_types else 'لا توجد ورديات',
            "عدد التأخيرات": delays_count,
            "تفاصيل التأخيرات": delay_details,
            "أيام الغياب": len(absent_days),
            "تفاصيل أيام الغياب": absent_days_formatted,
            "الدوام الإضافي": overtime_count,
            "تفاصيل الورديات الصباحية": morning_shift_details,
            "تفاصيل الخروج المتأخر": late_departure_details,
            "تفاصيل الورديات المزدوجة": double_shift_details
        })

    return pd.DataFrame(summary)

def export_to_word(summary_df):
    doc = Document()
    doc.add_heading("تقرير الحضور والانصراف العام", level=0)
    doc.add_paragraph("---")
    for index, row in summary_df.iterrows():
        doc.add_heading(f"بيانات الموظف: {row['الإسم']}", level=1)
        summary_paragraph = doc.add_paragraph()
        summary_paragraph.add_run(
            f"إجمالي الورديات: {row['عدد الورديات']}, "
            f"أنواعها: {row['نوع الورديات']}, "
            f"تأخيرات: {row['عدد التأخيرات']}, "
            f"غياب: {row['أيام الغياب']}, "
            f"دوام إضافي: {row['الدوام الإضافي']} وردية"
        )
        doc.add_paragraph()

        if row["تفاصيل الورديات المزدوجة"]:
            doc.add_paragraph("تفاصيل الورديات المزدوجة:")
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "التاريخ"
            hdr_cells[1].text = "الحضور"
            hdr_cells[2].text = "الانصراف"
            hdr_cells[3].text = "المدة"
            for d in row["تفاصيل الورديات المزدوجة"]:
                table.add_row().cells[:] = d
            doc.add_paragraph()

        if row["تفاصيل الورديات الصباحية"]:
            doc.add_paragraph("الورديات الصباحية:")
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "التاريخ"
            hdr_cells[1].text = "الحضور"
            hdr_cells[2].text = "الانصراف"
            for d in row["تفاصيل الورديات الصباحية"]:
                table.add_row().cells[:] = d
            doc.add_paragraph()

        if row["تفاصيل التأخيرات"]:
            doc.add_paragraph("تفاصيل التأخيرات:")
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "التاريخ"
            hdr_cells[1].text = "وقت التأخير"
            for d in row["تفاصيل التأخيرات"]:
                table.add_row().cells[:] = d
            doc.add_paragraph()

        if row["تفاصيل الخروج المتأخر"]:
            doc.add_paragraph("الخروج المتأخر:")
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "التاريخ"
            hdr_cells[1].text = "وقت الخروج"
            for d in row["تفاصيل الخروج المتأخر"]:
                table.add_row().cells[:] = d
            doc.add_paragraph()

        if row["تفاصيل أيام الغياب"]:
            doc.add_paragraph("تفاصيل أيام الغياب:")
            table = doc.add_table(rows=1, cols=1)
            table.style = 'Table Grid'
            table.rows[0].cells[0].text = "التاريخ (اليوم)"
            for d in row["تفاصيل أيام الغياب"]:
                table.add_row().cells[0].text = d[0]
            doc.add_paragraph()

    output = BytesIO()
    doc.save(output)
    return output

# Streamlit app interface
st.title("تحليل الحضور وإنشاء تقرير Word")
uploaded_file = st.file_uploader("ارفع ملف PDF", type=["pdf"])
if uploaded_file is not None:
    raw_data = load_pdf_data(uploaded_file)
    df_summary = process_attendance_data(raw_data)
    if not df_summary.empty:
        word_output = export_to_word(df_summary)
        st.success("تم توليد التقرير بنجاح!")
        st.download_button("📄 تحميل تقرير Word", word_output.getvalue(), file_name="تقرير_الحضور.docx")
    else:
        st.warning("لم يتم استخراج بيانات صالحة من الملف.")
